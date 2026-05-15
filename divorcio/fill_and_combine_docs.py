#!/usr/bin/env python3
"""
Script to fill template variables in divorce document templates and combine them into one document.
"""

import argparse
import json
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

# Keys only for JSON sidecars (never appear in Word {{...}} templates)
_META_VARIABLE_KEYS = frozenset({"_comment", "_note"})


def bracketed_missing_placeholder(var_name: str) -> str:
    """Visible placeholder when a merge field has no value (empty, whitespace, or absent from the JSON)."""
    return f"[MISSING — {var_name}]"


def resolve_merge_value(var_name: str, variables: dict) -> str:
    """
    Return the string to substitute for {{var_name}}.
    Missing or empty values become [MISSING — var_name] so reviewers can search the packet.
    """
    if var_name in _META_VARIABLE_KEYS:
        return bracketed_missing_placeholder(var_name)
    if var_name not in variables:
        return bracketed_missing_placeholder(var_name)
    raw = variables[var_name]
    if raw is None:
        return bracketed_missing_placeholder(var_name)
    s = str(raw).strip()
    if not s:
        return bracketed_missing_placeholder(var_name)
    return str(raw)


def replace_variables(text, variables):
    """Replace {{variable}} placeholders with values, or [MISSING — name] when unset."""
    if not text:
        return text

    pattern = r"\{\{([a-zA-Z0-9_]+)\}\}"

    def replacer(match):
        return resolve_merge_value(match.group(1), variables)

    return re.sub(pattern, replacer, text)


def replace_variables_in_paragraph(paragraph, variables):
    """Replace variables in a paragraph; unset fields show as [MISSING — var]."""
    full_text = paragraph.text
    if not full_text or "{{" not in full_text:
        return

    replaced_text = replace_variables(full_text, variables)
    if replaced_text == full_text:
        return

    paragraph.clear()
    run = paragraph.add_run(replaced_text)
    run.font.size = Pt(11)


def process_document(doc_path, variables):
    """Process a single document, replacing all variables."""
    try:
        doc = Document(doc_path)
        
        # Replace variables in paragraphs
        for paragraph in doc.paragraphs:
            replace_variables_in_paragraph(paragraph, variables)
        
        # Replace variables in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_variables_in_paragraph(paragraph, variables)
        
        # Replace variables in headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                for paragraph in header.paragraphs:
                    replace_variables_in_paragraph(paragraph, variables)
            
            for footer in [section.footer, section.first_page_footer]:
                for paragraph in footer.paragraphs:
                    replace_variables_in_paragraph(paragraph, variables)
        
        return doc
    except Exception as e:
        print(f"Error processing {doc_path}: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return None


def combine_documents(documents, output_path):
    """Combine multiple documents into one."""
    if not documents:
        print("No documents to combine", file=sys.stderr)
        return False
    
    # Start with the first document
    combined_doc = documents[0]
    
    # Add a page break and append each subsequent document
    for doc in documents[1:]:
        # Add page break before next document
        combined_doc.add_page_break()
        
        # Copy paragraphs from the document
        for paragraph in doc.paragraphs:
            new_paragraph = combined_doc.add_paragraph()
            # Copy paragraph properties
            new_paragraph.style = paragraph.style
            new_paragraph.alignment = paragraph.alignment
            
            # Copy runs with formatting
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                if run.bold is not None:
                    new_run.bold = run.bold
                if run.italic is not None:
                    new_run.italic = run.italic
                if run.underline is not None:
                    new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with same dimensions
            new_table = combined_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style
            
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_cell = new_table.rows[i].cells[j]
                    # Copy cell content
                    for paragraph in cell.paragraphs:
                        new_para = new_cell.add_paragraph()
                        new_para.style = paragraph.style
                        for run in paragraph.runs:
                            new_run = new_para.add_run(run.text)
                            if run.bold is not None:
                                new_run.bold = run.bold
                            if run.italic is not None:
                                new_run.italic = run.italic
    
    # Save the combined document
    combined_doc.save(output_path)
    return True


def get_docx_files(directory, exclude_output=None):
    """Get all .docx files in the directory, excluding the key file and output files."""
    doc_dir = Path(directory)
    docx_files = []
    
    for file_path in doc_dir.glob("*.docx"):
        file_name_upper = file_path.name.upper()
        
        # Skip excluded files:
        # - Files with "KEY" in name (the variables key file)
        # - Files starting with "test" (test outputs)
        # - Files containing "test_combined" or "combined" and "test" (test outputs)
        # - The output file itself (if specified)
        if ("KEY" in file_name_upper or 
            file_name_upper.startswith("TEST") or
            ("COMBINED" in file_name_upper and "TEST" in file_name_upper)):
            continue
        
        # Skip the output file if it exists
        if exclude_output and file_path.name == exclude_output:
            continue
            
        docx_files.append(file_path)
    
    # Sort for consistent ordering
    return sorted(docx_files)


def parse_variables_from_args(args_dict):
    """Parse variables from argparse namespace."""
    variables = {}
    
    # Get all variables from args, excluding special args
    exclude_keys = {'variables_file', 'output', 'directory', 'verbose'}
    
    for key, value in args_dict.items():
        if key not in exclude_keys and value is not None:
            variables[key] = value
    
    return variables


def load_variables_from_file(file_path):
    """Load variables from a JSON file."""
    try:
        with open(file_path, 'r') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading variables file {file_path}: {e}", file=sys.stderr)
        return None


def main():
    parser = argparse.ArgumentParser(
        description='Fill template variables in divorce documents and combine them into one document.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Using CLI arguments:
  python fill_and_combine_docs.py --plaintiff "John Doe" --defendant "Jane Doe" --county "New York"
  
  # Using a JSON file:
  python fill_and_combine_docs.py --variables-file variables.json
  
  # Combine both:
  python fill_and_combine_docs.py --variables-file variables.json --plaintiff "John Doe"
        """
    )
    
    # Common variables
    parser.add_argument('--plaintiff', help='Full name of plaintiff')
    parser.add_argument('--defendant', help='Full name of defendant')
    parser.add_argument('--county', help='County of jurisdiction')
    parser.add_argument('--venue-basis', dest='venue_basis', help='How county was chosen, which section of CPLR')
    parser.add_argument('--p-address', dest='p_address', help='Complete address of plaintiff')
    parser.add_argument('--d-address', dest='d_address', help='Complete address of defendant')
    parser.add_argument('--p-or-d-residence', dest='p_or_d_residence', help='Choose P/D address which forms basis of venue')
    parser.add_argument('--address', help='Insert P or D address based on previous answer')
    parser.add_argument('--todays-date', dest='todays_date', help="Today's date")
    parser.add_argument('--grounds', help='Grounds for divorce - cite DRL')
    parser.add_argument('--d-ssn', dest='d_ssn', help='Defendant SSN')
    parser.add_argument('--p-ssn', dest='p_ssn', help='Plaintiff SSN')
    parser.add_argument('--religious-or-not-religious', dest='religious_or_not_religious', 
                       help='Was the marriage performed by clergy?')
    parser.add_argument('--p-dob', dest='p_dob', help='Plaintiff date of birth')
    parser.add_argument('--d-dob', dest='d_dob', help='Defendant date of birth')
    parser.add_argument('--marriage-date', dest='marriage_date', help='Date of marriage')
    parser.add_argument('--marriage-city', dest='marriage_city', help='City of marriage')
    parser.add_argument('--marriage-county', dest='marriage_county', help='County of marriage')
    parser.add_argument('--marriage-state', dest='marriage_state', help='State of marriage')
    parser.add_argument('--plaintiff-or-defendant', dest='plaintiff_or_defendant', 
                       help='Plaintiff or Defendant')
    parser.add_argument('--maiden-name', dest='maiden_name', help='Maiden name')
    parser.add_argument('--attorney-name', dest='attorney_name', help='Name & address & phone of attorney')
    parser.add_argument('--atty-state', dest='atty_state', help="Lawyer's state")
    parser.add_argument('--atty-county', dest='atty_county', help="Lawyer's county")
    parser.add_argument('--service-inside-outside-ny', dest='service_inside_outside_ny', 
                       help='Where was D served? Inside or outside NY')
    parser.add_argument('--d-gender', dest='d_gender', help='Defendant gender')
    parser.add_argument('--p-gender', dest='p_gender', help='Plaintiff gender')
    parser.add_argument('--filed-date', dest='filed_date', help='Date summons with notice was filed')
    parser.add_argument('--service-date', dest='service_date', help='Date SWN was served on D')
    parser.add_argument('--index', help='Index number')
    parser.add_argument('--separation-date', dest='separation_date', help='Separation date')
    parser.add_argument('--judgment-date', dest='judgment_date', help='Date signed by judge')
    parser.add_argument('--entry-date', dest='entry_date', help='Date entered')
    
    # File options
    parser.add_argument('--variables-file', dest='variables_file', 
                       help='JSON file containing variables (keys should match variable names without {{}})')
    parser.add_argument('--output', '-o', default='combined_divorce_docs.docx',
                       help='Output filename (default: combined_divorce_docs.docx)')
    parser.add_argument('--directory', '-d', default='.',
                       help='Directory containing docx files (default: current directory)')
    parser.add_argument('--verbose', '-v', action='store_true',
                       help='Verbose output')
    
    args = parser.parse_args()
    
    # Load variables
    variables = {}
    
    # Load from file if provided
    if args.variables_file:
        file_vars = load_variables_from_file(args.variables_file)
        if file_vars is None:
            return 1
        variables.update(file_vars)
    
    # Add variables from CLI args (will override file variables)
    cli_vars = parse_variables_from_args(vars(args))
    variables.update(cli_vars)
    
    if not variables:
        print("Error: No variables provided. Use --variables-file or provide variables via CLI arguments.", 
              file=sys.stderr)
        return 1
    
    if args.verbose:
        print(f"Loaded {len(variables)} variables", file=sys.stderr)
        print(f"Variables: {list(variables.keys())}", file=sys.stderr)
    
    # Get all docx files (exclude the output file if it exists)
    docx_files = get_docx_files(args.directory, exclude_output=args.output)
    
    if not docx_files:
        print(f"No .docx files found in {args.directory}", file=sys.stderr)
        return 1
    
    if args.verbose:
        print(f"Found {len(docx_files)} documents to process:", file=sys.stderr)
        for f in docx_files:
            print(f"  - {f.name}", file=sys.stderr)
    
    # Process each document
    processed_docs = []
    for doc_path in docx_files:
        if args.verbose:
            print(f"Processing {doc_path.name}...", file=sys.stderr)
        
        doc = process_document(doc_path, variables)
        if doc:
            processed_docs.append(doc)
        else:
            print(f"Warning: Failed to process {doc_path.name}", file=sys.stderr)
    
    if not processed_docs:
        print("Error: No documents were successfully processed", file=sys.stderr)
        return 1
    
    # Combine documents
    output_path = Path(args.directory) / args.output
    if args.verbose:
        print(f"Combining {len(processed_docs)} documents into {output_path}...", file=sys.stderr)
    
    if combine_documents(processed_docs, output_path):
        print(f"Successfully created combined document: {output_path}")
        return 0
    else:
        print("Error: Failed to combine documents", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
